import json
import random
import time
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from fastapi import HTTPException, UploadFile
from fastapi.responses import StreamingResponse

from app.services.gemini_service import generate_json


# ============================================================
# DOCX HELPERS
# ============================================================

def collect_docx_text_locations(document: Document):

    locations = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):

        text = paragraph.text.strip()

        if text:
            locations.append({
                "location_id": f"paragraph:{paragraph_index}",
                "kind": "paragraph",
                "text": text,
                "paragraph_index": paragraph_index,
            })

    for table_index, table in enumerate(document.tables):

        for row_index, row in enumerate(table.rows):

            for cell_index, cell in enumerate(row.cells):

                for paragraph_index, paragraph in enumerate(cell.paragraphs):

                    text = paragraph.text.strip()

                    if text:
                        locations.append({
                            "location_id": (
                                f"table:{table_index}:"
                                f"row:{row_index}:"
                                f"cell:{cell_index}:"
                                f"paragraph:{paragraph_index}"
                            ),
                            "kind": "table_paragraph",
                            "text": text,
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index,
                        })

    return locations


def replace_paragraph_text_preserving_style(
    paragraph,
    new_text: str
):

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


def get_docx_paragraph_by_location(
    document: Document,
    location: dict
):

    if location["kind"] == "paragraph":

        return document.paragraphs[
            location["paragraph_index"]
        ]

    if location["kind"] == "table_paragraph":

        table = document.tables[
            location["table_index"]
        ]

        row = table.rows[
            location["row_index"]
        ]

        cell = row.cells[
            location["cell_index"]
        ]

        return cell.paragraphs[
            location["paragraph_index"]
        ]

    raise ValueError(
        "Unsupported DOCX text location."
    )


# ============================================================
# DOCX OPTIMIZER
# ============================================================

async def optimize_docx(
    file: UploadFile,
    job_description: str
):

    filename = file.filename or ""

    if not filename.lower().endswith(".docx"):

        raise HTTPException(
            status_code=400,
            detail="Please upload a DOCX file."
        )

    if not job_description.strip():

        raise HTTPException(
            status_code=400,
            detail="Job description is required."
        )

    file_content = await file.read()

    if not file_content:

        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty."
        )

    try:

        document = Document(
            BytesIO(file_content)
        )

        locations = collect_docx_text_locations(
            document
        )

        editable_locations = []

        current_section = ""

        for location in locations:

            text = location["text"].strip()

            if not text:
                continue

            upper = text.upper()

            if upper in {
                "SUMMARY",
                "EDUCATION",
                "PROJECTS",
                "SKILLS",
                "CERTIFICATIONS",
                "CODING PROFILES",
            }:

                current_section = upper

                continue

            if current_section not in {
                "SUMMARY",
                "PROJECTS",
                "SKILLS",
            }:

                continue

            if len(text) < 20:
                continue

            editable_locations.append({
                "location_id":
                    location["location_id"],

                "section":
                    current_section,

                "text":
                    text,

                "max_characters":
                    int(len(text) * 1.25),
            })

        if not editable_locations:

            raise HTTPException(
                status_code=400,
                detail=(
                    "No editable resume sections "
                    "were found in the DOCX."
                )
            )

        prompt = f"""
You are an expert ATS resume editor.

Improve the resume text for the supplied job description.

STRICT RULES:

- Preserve every existing fact.
- Do not invent skills.
- Do not invent achievements.
- Do not invent metrics.
- Do not invent projects.
- Do not invent experience.
- Do not invent certifications.
- Keep the same meaning.
- Improve clarity and impact.
- Improve action verbs.
- Improve job relevance.
- Keep each improved text close to original length.
- Do not change headings.
- Do not add sections.
- If text is already strong, keep it nearly unchanged.
- Return ONLY valid JSON.

JOB DESCRIPTION:

{job_description[:6000]}

TEXT LOCATIONS:

{json.dumps(
    editable_locations,
    ensure_ascii=False
)}

Return exactly:

{{
    "changes": [
        {{
            "location_id": "paragraph:5",
            "original": "original text",
            "improved": "improved text"
        }}
    ]
}}
"""

        result = generate_json(prompt)

        changes = result.get(
            "changes",
            []
        )

        location_map = {
            location["location_id"]:
                location

            for location in locations
        }

        applied_changes = []

        for change in changes:

            location_id = change.get(
                "location_id",
                ""
            )

            improved_text = change.get(
                "improved",
                ""
            ).strip()

            if not location_id:
                continue

            if not improved_text:
                continue

            location = location_map.get(
                location_id
            )

            if not location:
                continue

            original_text = location[
                "text"
            ].strip()

            if original_text != change.get(
                "original",
                ""
            ).strip():

                continue

            paragraph = (
                get_docx_paragraph_by_location(
                    document,
                    location
                )
            )

            replace_paragraph_text_preserving_style(
                paragraph,
                improved_text
            )

            applied_changes.append({
                "location_id":
                    location_id,

                "original":
                    original_text,

                "improved":
                    improved_text,
            })

        output = BytesIO()

        document.save(output)

        output.seek(0)

        original_name = Path(
            filename
        ).stem

        output_name = (
            f"{original_name}_optimized.docx"
        )

        return StreamingResponse(
            output,
            media_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            headers={
                "Content-Disposition":
                    f'attachment; filename="{output_name}"',

                "X-Applied-Changes":
                    str(len(applied_changes)),
            }
        )

    except HTTPException:
        raise

    except Exception as error:

        print(
            "DOCX optimization error:",
            error
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Could not optimize DOCX: {str(error)}"
            )
        )


# ============================================================
# PDF HELPERS
# ============================================================

def clean_text(value: str) -> str:

    return " ".join(
        str(value).split()
    ).strip()


def normalize_heading(value: str) -> str:

    normalized = clean_text(
        value
    ).upper()

    normalized = normalized.rstrip(
        ":|-–—"
    ).strip()

    normalized = normalized.replace(
        " ",
        ""
    )

    return normalized


def replace_unsupported_characters(
    value: str
) -> str:

    replacements = {
        "•": "-",
        "●": "-",
        "▪": "-",
        "◦": "-",
        "–": "-",
        "—": "-",
        "−": "-",
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',
        "\u00a0": " ",
    }

    for old_character, new_character in replacements.items():

        value = value.replace(
            old_character,
            new_character
        )

    return value


def safe_pdf_text(value: str) -> str:

    value = replace_unsupported_characters(
        str(value)
    )

    return clean_text(value)


def format_skills_text(value: str) -> str:

    value = replace_unsupported_characters(
        str(value)
    )

    value = clean_text(value)

    value = value.replace(
        "EST APIs",
        "REST APIs"
    )

    skill_labels = [
        "Languages:",
        "Web Technologies:",
        "Tools & Platforms:",
        "Database:",
        "AI/ML:",
        "Cloud & APIs:",
    ]

    for label in skill_labels:

        value = value.replace(
            label,
            "\n" + label
        )

    lines = []

    for line in value.splitlines():

        cleaned_line = clean_text(
            line
        )

        if cleaned_line:
            lines.append(
                cleaned_line
            )

    return "\n".join(lines).strip()


def shorten_at_word(
    value: str,
    maximum_length: int
) -> str:

    if len(value) <= maximum_length:
        return value

    shortened = value[
        :maximum_length
    ]

    if " " in shortened:

        shortened = shortened.rsplit(
            " ",
            1
        )[0]

    return shortened.rstrip(
        " ,;:-"
    )


# ============================================================
# PDF OPTIMIZER
# ============================================================

async def optimize_pdf(
    file: UploadFile,
    job_description: str
):

    if not file.filename:

        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF file."
        )

    if not file.filename.lower().endswith(
        ".pdf"
    ):

        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF file."
        )

    if not job_description.strip():

        raise HTTPException(
            status_code=400,
            detail="Please provide a job description."
        )

    pdf_bytes = await file.read()

    if not pdf_bytes:

        raise HTTPException(
            status_code=400,
            detail="The uploaded PDF is empty."
        )

    document = None

    try:

        document = fitz.open(
            stream=pdf_bytes,
            filetype="pdf"
        )

        heading_aliases = {

            "SUMMARY": "SUMMARY",
            "PROFILE": "SUMMARY",
            "PROFESSIONALSUMMARY": "SUMMARY",
            "CAREERSUMMARY": "SUMMARY",
            "OBJECTIVE": "SUMMARY",
            "CAREEROBJECTIVE": "SUMMARY",
            "ABOUTME": "SUMMARY",

            "PROJECTS": "PROJECTS",
            "PROJECT": "PROJECTS",
            "PERSONALPROJECTS": "PROJECTS",
            "ACADEMICPROJECTS": "PROJECTS",
            "KEYPROJECTS": "PROJECTS",
            "PROJECTEXPERIENCE": "PROJECTS",

            "SKILLS": "SKILLS",
            "TECHNICALSKILLS": "SKILLS",
            "CORESKILLS": "SKILLS",
            "KEYSKILLS": "SKILLS",
            "SKILLSET": "SKILLS",
            "TECHNOLOGIES": "SKILLS",
            "TECHNOLOGYSTACK": "SKILLS",

            "EDUCATION": "EDUCATION",
            "ACADEMICBACKGROUND": "EDUCATION",
            "ACADEMICQUALIFICATIONS":
                "EDUCATION",

            "EXPERIENCE": "EXPERIENCE",
            "WORKEXPERIENCE":
                "EXPERIENCE",
            "PROFESSIONALEXPERIENCE":
                "EXPERIENCE",
            "INTERNSHIP": "EXPERIENCE",
            "INTERNSHIPS": "EXPERIENCE",

            "CERTIFICATIONS":
                "CERTIFICATIONS",
            "CERTIFICATES":
                "CERTIFICATIONS",
            "COURSES":
                "CERTIFICATIONS",

            "CODINGPROFILES":
                "CODING PROFILES",
            "CODINGPROFILE":
                "CODING PROFILES",
            "PROFILES":
                "CODING PROFILES",

            "ACHIEVEMENTS":
                "ACHIEVEMENTS",
            "AWARDS":
                "ACHIEVEMENTS",

            "POSITIONSOFRESPONSIBILITY":
                "POSITIONS OF RESPONSIBILITY",

            "CONTACT":
                "CONTACT",

            "CONTACTDETAILS":
                "CONTACT",
        }

        allowed_sections = {
            "SUMMARY",
            "PROJECTS",
            "SKILLS",
        }

        skill_labels = [
            "Languages:",
            "Web Technologies:",
            "Tools & Platforms:",
            "Database:",
            "AI/ML:",
            "Cloud & APIs:",
        ]

        def detect_heading(value):

            normalized = normalize_heading(
                value
            )

            return heading_aliases.get(
                normalized
            )

        editable_blocks = []

        detected_headings = []

        current_section = ""

        for page_number, page in enumerate(
            document
        ):

            page_dict = page.get_text(
                "dict"
            )

            for block_index, block in enumerate(
                page_dict.get(
                    "blocks",
                    []
                )
            ):

                lines = block.get(
                    "lines",
                    []
                )

                if not lines:
                    continue

                block_spans = []
                block_lines = []

                for line in lines:

                    spans = line.get(
                        "spans",
                        []
                    )

                    if not spans:
                        continue

                    block_spans.extend(
                        spans
                    )

                    line_text = clean_text(
                        " ".join(
                            span.get(
                                "text",
                                ""
                            )

                            for span in spans

                            if span.get(
                                "text",
                                ""
                            ).strip()
                        )
                    )

                    if line_text:
                        block_lines.append(
                            line_text
                        )

                if (
                    not block_lines
                    or not block_spans
                ):
                    continue

                block_text = clean_text(
                    " ".join(
                        block_lines
                    )
                )

                if not block_text:
                    continue

                matched_section = detect_heading(
                    block_text
                )

                if matched_section:

                    current_section = (
                        matched_section
                    )

                    if (
                        matched_section
                        not in detected_headings
                    ):

                        detected_headings.append(
                            matched_section
                        )

                    continue

                if current_section not in allowed_sections:
                    continue

                should_edit = False

                if current_section == "SUMMARY":

                    should_edit = True

                elif current_section == "SKILLS":

                    should_edit = True

                elif current_section == "PROJECTS":

                    first_character = (
                        block_text.lstrip()[:1]
                    )

                    bullet_characters = {
                        "•",
                        "●",
                        "▪",
                        "◦",
                        "-",
                        "–",
                        "—",
                    }

                    if (
                        first_character
                        in bullet_characters
                    ):

                        should_edit = True

                if not should_edit:
                    continue

                x0, y0, x1, y1 = (
                    block["bbox"]
                )

                first_span = block_spans[0]

                editable_blocks.append({
                    "block_id": (
                        f"page:{page_number}:"
                        f"block:{block_index}"
                    ),

                    "page_number":
                        page_number,

                    "section":
                        current_section,

                    "text":
                        block_text,

                    "x0":
                        float(x0),

                    "y0":
                        float(y0),

                    "x1":
                        float(x1),

                    "y1":
                        float(y1),

                    "font_size":
                        float(
                            first_span.get(
                                "size",
                                10
                            )
                        ),

                    "max_characters":
                        max(
                            len(block_text),
                            20
                        ),
                })

        if not editable_blocks:

            detected_text = (
                ", ".join(
                    detected_headings
                )
                if detected_headings
                else "None"
            )

            raise HTTPException(
                status_code=400,
                detail=(
                    "No editable resume sections "
                    "were found. "
                    f"Detected headings: "
                    f"{detected_text}."
                )
            )

        prompt = f"""
You are optimizing selected resume content
for the supplied job description.

JOB DESCRIPTION:

{job_description[:6000]}

RESUME BLOCKS:

{json.dumps(
    editable_blocks,
    ensure_ascii=False
)}

Rules:

1. Preserve every candidate fact.
2. Do not invent experience, skills, tools,
   technologies, responsibilities, metrics,
   qualifications or achievements.
3. Do not change project names.
4. Do not change company names.
5. Do not change job titles or project roles.
6. Do not change dates, numbers, URLs or certifications.
7. Do not add skills absent from the original resume block.
8. Improve SUMMARY for job alignment.
9. Improve PROJECT descriptions.
10. Reorder SKILLS using only existing skills.
11. Preserve skill category labels.
12. Keep skill categories separate.
13. Preserve original meaning.
14. Do not exceed max_characters.
15. Keep block_id unchanged.
16. Return one result for every supplied block.
17. Return valid JSON only.
18. Do not use Markdown.
19. Do not provide explanations.
20. Use ASCII hyphens.
21. Do not leave incomplete sentences.
22. Project descriptions must have a leading hyphen.
23. Do not rename Website.

Return exactly:

{{
    "changes": [
        {{
            "block_id": "page:0:block:4",
            "improved": "Improved text"
        }}
    ]
}}
"""

        result = generate_json(prompt)

        changes = result.get(
            "changes",
            []
        )

        if not isinstance(
            changes,
            list
        ):

            raise HTTPException(
                status_code=500,
                detail=(
                    "Gemini returned an invalid "
                    "changes format."
                )
            )

        blocks_by_id = {
            block["block_id"]:
                block

            for block in editable_blocks
        }

        valid_changes = []

        for change in changes:

            if not isinstance(
                change,
                dict
            ):
                continue

            block_id = change.get(
                "block_id"
            )

            original_block = (
                blocks_by_id.get(
                    block_id
                )
            )

            if not original_block:
                continue

            raw_improved_text = change.get(
                "improved",
                ""
            )

            if (
                original_block["section"]
                == "SKILLS"
            ):

                improved_text = (
                    format_skills_text(
                        raw_improved_text
                    )
                )

            else:

                improved_text = (
                    safe_pdf_text(
                        raw_improved_text
                    )
                )

            if not improved_text:
                continue

            max_characters = (
                original_block[
                    "max_characters"
                ]
            )

            if (
                original_block["section"]
                != "SKILLS"
            ):

                improved_text = (
                    shorten_at_word(
                        improved_text,
                        max_characters
                    )
                )

            if (
                original_block["section"]
                == "PROJECTS"
            ):

                improved_text = (
                    improved_text.lstrip(
                        "- "
                    )
                )

                improved_text = (
                    "- "
                    + improved_text
                )

                improved_text = (
                    shorten_at_word(
                        improved_text,
                        max_characters
                    )
                )

            valid_changes.append({
                "block_id":
                    block_id,

                "improved":
                    improved_text,
            })

        if not valid_changes:

            raise HTTPException(
                status_code=500,
                detail=(
                    "Gemini did not return any "
                    "valid resume changes."
                )
            )

        # Remove original text

        for change in valid_changes:

            original_block = (
                blocks_by_id[
                    change["block_id"]
                ]
            )

            page = document[
                original_block["page_number"]
            ]

            original_rect = fitz.Rect(
                original_block["x0"] - 1,
                original_block["y0"] - 1,
                original_block["x1"] + 2,
                original_block["y1"] + 2,
            )

            page.add_redact_annot(
                original_rect,
                fill=(1, 1, 1)
            )

        for page in document:

            page.apply_redactions()

        # Insert optimized text

        for change in valid_changes:

            original_block = (
                blocks_by_id[
                    change["block_id"]
                ]
            )

            improved_text = (
                change["improved"]
            )

            page = document[
                original_block["page_number"]
            ]

            original_font_size = (
                original_block[
                    "font_size"
                ]
            )

            page_width = page.rect.width

            if (
                original_block["section"]
                == "SUMMARY"
            ):

                extra_height = 14
                extra_width = 80

            elif (
                original_block["section"]
                == "SKILLS"
            ):

                extra_height = 8
                extra_width = 40

            else:

                extra_height = 8
                extra_width = 40

            replacement_rect = fitz.Rect(
                original_block["x0"],
                original_block["y0"] - 1,
                min(
                    original_block["x1"]
                    + extra_width,
                    page_width - 20
                ),
                original_block["y1"]
                + extra_height,
            )

            font_sizes = [
                original_font_size,
                original_font_size - 0.5,
                original_font_size - 1,
                original_font_size - 1.5,
                original_font_size - 2,
            ]

            inserted = False

            for font_size in font_sizes:

                safe_font_size = max(
                    float(font_size),
                    7.0
                )

                insert_result = (
                    page.insert_textbox(
                        replacement_rect,
                        improved_text,
                        fontsize=safe_font_size,
                        fontname="helv",
                        color=(0, 0, 0),
                        align=fitz.TEXT_ALIGN_LEFT,
                    )
                )

                if insert_result >= 0:

                    inserted = True

                    break

            if not inserted:

                larger_rect = fitz.Rect(
                    replacement_rect.x0,
                    replacement_rect.y0,
                    replacement_rect.x1,
                    replacement_rect.y1 + 15,
                )

                page.insert_textbox(
                    larger_rect,
                    improved_text,
                    fontsize=7.0,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

        output = BytesIO()

        document.save(
            output,
            garbage=4,
            deflate=True
        )

        output.seek(0)

        document.close()

        document = None

        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={
                "Content-Disposition":
                    'attachment; filename="optimized_resume.pdf"'
            }
        )

    except HTTPException:

        if document is not None:
            document.close()

        raise

    except Exception as error:

        if document is not None:
            document.close()

        print(
            "PDF optimization error:",
            error
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Could not optimize PDF: "
                f"{str(error)}"
            )
        )
    