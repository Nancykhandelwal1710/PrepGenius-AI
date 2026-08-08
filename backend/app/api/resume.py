from io import BytesIO

import pdfplumber
from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile

from app.models import ResumeRequest
from app.services.resume_service import resume_suggestions


router = APIRouter(tags=["Resume"])


@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    filename = (file.filename or "").lower()

    # -------------------------
    # PDF
    # -------------------------
    if filename.endswith(".pdf"):
        try:
            text_parts = []

            with pdfplumber.open(file.file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(page_text)

            text = "\n".join(text_parts).strip()

            if not text:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "No readable text was found in this PDF. "
                        "The PDF may be scanned or image-based."
                    ),
                )

            return {"text": text}

        except HTTPException:
            raise

        except Exception as e:
            print("PDF extraction failed:", repr(e))

            raise HTTPException(
                status_code=500,
                detail=f"Could not extract text from PDF: {str(e)}",
            )

    # -------------------------
    # DOCX
    # -------------------------
    if filename.endswith(".docx"):
        try:
            contents = await file.read()

            document = Document(
                BytesIO(contents)
            )

            text_parts = []

            # Normal paragraphs
            for paragraph in document.paragraphs:
                paragraph_text = paragraph.text.strip()

                if paragraph_text:
                    text_parts.append(
                        paragraph_text
                    )

            # Tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = (
                                paragraph.text.strip()
                            )

                            if paragraph_text:
                                text_parts.append(
                                    paragraph_text
                                )

            text = "\n".join(
                text_parts
            ).strip()

            if not text:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "No readable text was found "
                        "in this DOCX file."
                    ),
                )

            return {"text": text}

        except HTTPException:
            raise

        except Exception as e:
            print(
                "DOCX extraction failed:",
                repr(e)
            )

            raise HTTPException(
                status_code=500,
                detail=(
                    "Could not extract text from DOCX: "
                    f"{str(e)}"
                ),
            )

    # -------------------------
    # Unsupported file
    # -------------------------
    raise HTTPException(
        status_code=400,
        detail="Only PDF and DOCX files are supported.",
    )


@router.post("/resume-suggestions")
def suggestions(data: ResumeRequest):
    return resume_suggestions(data)
