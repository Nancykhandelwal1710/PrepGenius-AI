from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pdfplumber
import os
import json
import re
import time
import random
import fitz
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from typing import Annotated
from docx import Document
from fastapi import File, Form, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from google import genai

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
print("Gemini Key:", GEMINI_API_KEY)

client = genai.Client(api_key=GEMINI_API_KEY)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
    "http://localhost:5173",
    "http://localhost:5174",
    "https://prepgenius-ai-career.vercel.app",
],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def home():
    return {"message": "Backend Running"}


@app.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    filename = file.filename.lower()

    # ---------- PDF ----------
    if filename.endswith(".pdf"):
        text = ""

        with pdfplumber.open(file.file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

        return {"text": text} 

    # ---------- DOCX ----------
    elif filename.endswith(".docx"):
        contents = await file.read()

        document = Document(BytesIO(contents))

        text_parts = []

        # Read normal paragraphs
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                text_parts.append(paragraph_text)

        # Read text inside tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text.strip()

                        if paragraph_text:
                            text_parts.append(paragraph_text)

        text = "\n".join(text_parts)

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No readable text was found in this DOCX file."
            )

        return {"text": text}

    raise HTTPException(
        status_code=400,
        detail="Only PDF and DOCX files are supported."
    )


class ATSRequest(BaseModel):
    resume_text: str
    job_description: str

class ResumeTailorRequest(BaseModel):
    resume_text: str
    job_description: str

class InterviewRequest(BaseModel):
    role: str
    resume_text: str = ""
    job_description: str = ""

class ResumeHealthRequest(BaseModel):
    resume_text: str
    ats_score: int = 0
    matched_skills: list[str] = []
    missing_skills: list[str] = []

def fallback_ats_analysis(resume_text, job_description):
    jd_words = re.findall(r"\b[a-zA-Z][a-zA-Z+#. ]{2,}\b", job_description.lower())
    resume_lower = resume_text.lower()

    stop_words = {
        "we", "are", "hiring", "with", "and", "the", "for", "this", "that",
        "candidate", "job", "role", "must", "have", "should", "will", "our",
        "your", "you", "experience", "knowledge", "skills", "ability"
    }

    possible_skills = []

    skill_phrases = [
        "recruitment", "onboarding", "employee engagement", "communication",
        "payroll coordination", "hr policy", "performance management",
        "documentation", "conflict resolution", "ms excel", "excel",
        "python", "java", "c++", "sql", "machine learning", "deep learning",
        "aws", "cloud", "react", "javascript", "html", "css", "sales",
        "marketing", "seo", "finance", "accounting", "teaching",
        "classroom management", "lesson planning", "content writing",
        "project management", "leadership", "customer service"
    ]

    for skill in skill_phrases:
        if skill in job_description.lower():
            possible_skills.append(skill.title())

    possible_skills = list(dict.fromkeys(possible_skills))

    matched = []
    missing = []

    for skill in possible_skills:
        if skill.lower() in resume_lower:
            matched.append(skill)
        else:
            missing.append(skill)

    score = 0
    if len(possible_skills) > 0:
        score = round((len(matched) / len(possible_skills)) * 100, 2)

    jd_lower = job_description.lower()

    if "hr" in jd_lower or "recruitment" in jd_lower or "onboarding" in jd_lower:
        domain = "Human Resources"
    elif "marketing" in jd_lower or "seo" in jd_lower:
        domain = "Marketing"
    elif "finance" in jd_lower or "accounting" in jd_lower:
        domain = "Finance"
    elif "teacher" in jd_lower or "teaching" in jd_lower:
        domain = "Teaching"
    elif "python" in jd_lower or "machine learning" in jd_lower or "developer" in jd_lower:
        domain = "Technology"
    else:
        domain = "General"

    if "intern" in jd_lower or "fresher" in jd_lower or "entry" in jd_lower:
        level = "Entry Level"
    elif "senior" in jd_lower or "manager" in jd_lower:
        level = "Senior Level"
    else:
        level = "Mid/General Level"

    return {
        "job_domain": domain,
        "experience_level": level,
        "ats_score": score,
        "required_skills": possible_skills,
        "matched_skills": matched,
        "missing_skills": missing
    }
def collect_docx_text_locations(document: Document):
    locations = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()

        if text:
            locations.append({
                "location_id": f"paragraph:{paragraph_index}",
                "kind": "paragraph",
                "text": text,
                "paragraph_index": paragraph_index,
            })

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    text = paragraph.text.strip()

                    if text:
                        locations.append({
                            "location_id": f"table:{table_index}:row:{row_index}:cell:{cell_index}:paragraph:{paragraph_index}",
                            "kind": "table_paragraph",
                            "text": text,
                            "table_index": table_index,
                            "row_index": row_index,
                            "cell_index": cell_index,
                            "paragraph_index": paragraph_index,
                        })

    return locations

def replace_paragraph_text_preserving_style(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


def get_docx_paragraph_by_location(document: Document, location: dict):
    if location["kind"] == "paragraph":
        return document.paragraphs[location["paragraph_index"]]

    if location["kind"] == "table_paragraph":
        table = document.tables[location["table_index"]]
        row = table.rows[location["row_index"]]
        cell = row.cells[location["cell_index"]]

        return cell.paragraphs[location["paragraph_index"]]

    raise ValueError("Unsupported DOCX text location.")

@app.post("/ats-score")
def ats_score(data: ATSRequest):
    print("ATS ROUTE HIT")
    print("NEW UNIVERSAL ATS FUNCTION RUNNING")

    try:
        prompt = f"""
You are an expert ATS (Applicant Tracking System) analyzer.

Your task is to compare ANY resume against ANY job description.

The job description may belong to any profession including:
Software Engineering, HR, Marketing, Finance, Teaching,
Healthcare, Sales, Operations, Design, Government,
Business, Law, Hospitality or any other industry.

Instructions:

1. Identify the job domain.

2. Identify the expected experience level.

3. Extract all important skills from ONLY the job description.

4. Compare those skills with the resume.

5. Return:
- required_skills
- matched_skills
- missing_skills

6. Calculate an ATS score (0-100).

Scoring should consider:
• Skill match
• Relevant experience
• ATS keywords
• Resume relevance

Do NOT invent skills.

Return ONLY valid JSON.

JOB DESCRIPTION:

{data.job_description[:6000]}

RESUME:

{data.resume_text[:6000]}

Return:

{{
    "job_domain":"",
    "experience_level":"",
    "ats_score":0,
    "required_skills":[],
    "matched_skills":[],
    "missing_skills":[]
}}
"""

        response = client.models.generate_content(
            model="gemini-3.5-flash",
            contents=prompt,
            config={
                "response_mime_type": "application/json"
            }
        )

        raw_text = response.text.strip()
        print("GEMINI RAW RESPONSE:", raw_text)

        raw_text = raw_text.replace("```json", "").replace("```", "").strip()

        json_match = re.search(r"\{.*\}", raw_text, re.DOTALL)

        if json_match:
            clean_json = json_match.group(0)
            result = json.loads(clean_json)
        else:
            raise Exception("Gemini did not return valid JSON")

        return {
            "job_domain": result.get("job_domain", "Unknown"),
            "experience_level": result.get("experience_level", "Unknown"),
            "ats_score": result.get("ats_score", 0),
            "required_skills": result.get("required_skills", []),
            "matched_skills": result.get("matched_skills", []),
            "missing_skills": result.get("missing_skills", [])
        }

    except Exception as e:
        print("ATS ERROR:", str(e))
        

        return {
            "job_domain": "Unknown",
            "experience_level": "Unknown",
            "ats_score": 0,
            "required_skills": [],
            "matched_skills": [],
            "missing_skills": [],
            "message": "AI analysis is temporarily unavailable. Please try again."
        }
    

@app.post("/suggestions")
def suggestions(data: ATSRequest):

    prompt = f"""
You are an expert ATS resume reviewer.

Resume:
{data.resume_text}

Job Description:
{data.job_description}

Generate resume improvement suggestions.

Rules:
- Suggestions must work for ANY profession.
- Do NOT assume the job is technical.
- Compare the resume against the job description.
- Mention missing skills only if they appear in the JD.
- Do not invent experience.
- Give 6-10 concise suggestions.
- Focus on:
  * Skills
  * Experience
  * Projects (if applicable)
  * Keywords
  * Achievements
  * Action verbs
  * Formatting
  * ATS optimization
- Return ONLY valid JSON.

Example:

{
  "suggestions":[
    "Highlight leadership experience.",
    "Include measurable achievements.",
    "Mention customer interaction experience.",
    "Use more ATS keywords from the job description."
  ]
}
"""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config={
                "response_mime_type": "application/json"
            }
        )

        result = json.loads(response.text)

        return {
            "suggestions": result.get("suggestions", [])
        }

    except Exception as e:

        return {
            "suggestions": [
                "Add more measurable achievements.",
                "Use action verbs to describe your work.",
                "Include keywords from the job description.",
                "Tailor your resume for each application."
            ]
        }


@app.post("/generate-interview-questions")
def generate_interview_questions(data: InterviewRequest):

    try:
        prompt = f"""
You are an interview coach.

Generate exactly 5 interview questions for this candidate.

Role:
{data.role}

Resume:
{data.resume_text[:2000]}

Job Description:
{data.job_description[:2000]}

Rules:
- Return only the questions.
- Number them from 1 to 5.
- Make questions practical and interview-style.
- Include technical, project-based, and HR questions.
"""

        response = client.models.generate_content(
            model="gemini-3.5-flash",
            contents=prompt
        )
        print("========== GEMINI QUESTION RESPONSE ==========")
        print(response.text)
        print("=============================================")
        questions_text = response.text

        questions = []
        for line in questions_text.split("\n"):
            line = line.strip()
            if line:
                line = (
                    line.replace("1.", "")
                    .replace("2.", "")
                    .replace("3.", "")
                    .replace("4.", "")
                    .replace("5.", "")
                    .strip()
                )
                questions.append(line)

        return {
            "questions": questions[:5]
        }

    except Exception as e:
        print("QUESTION GENERATION ERROR:", e)
        return {
            "questions": [],
            "error": str(e)
        }
    
class AnswerEvaluationRequest(BaseModel):
    question: str
    answer: str


@app.post("/evaluate-answer")
def evaluate_answer(data: AnswerEvaluationRequest):

    try:
        prompt = f"""
You are a Senior Technical Recruiter conducting a real interview.

Interview Question:
{data.question}

Candidate Answer:
{data.answer}

Evaluate exactly like a recruiter.

Scoring Rules:

Technical Accuracy (0-4)
Communication (0-2)
Completeness (0-2)
Confidence (0-1)
Practical Example (0-1)

Never give high marks for long but incorrect answers.

If the answer is factually wrong,
deduct marks heavily.

If the answer avoids the question,
give low score.

If the answer is copied,
reduce confidence.

After evaluating provide:

Return ONLY JSON in this exact format:

score (number)
feedback (string)
technical_accuracy (number)
completeness (number)
communication (number)
confidence (number)
practical_example (number)
conciseness (number)
strengths (array of strings)
weaknesses (array of strings)
improvements (array of strings)
verdict (string)
followup_question (string)
ideal_answer (string)

"""

        response = client.models.generate_content(
            model="gemini-3.5-flash",
            contents=prompt,
            config={
                "response_mime_type": "application/json"
            }
        )

        print("========== GEMINI RAW ==========")
        print(response.text)
        print("================================")

        clean_text = response.text.strip()
        # Remove markdown if Gemini adds it
        clean_text = clean_text.replace("```json", "").replace("```", "").strip()
        
        result = json.loads(clean_text)

        print("========== RESULT ==========")
        print(result)
        print("============================")

        return {
            "score": result.get("score", 0),

            "feedback": {
                "feedback": result.get("feedback", ""),
                "technical_accuracy": result.get("technical_accuracy", 0),
                "completeness": result.get("completeness", 0),
                "communication": result.get("communication", 0),
                "confidence": result.get("confidence", 0),
                "practical_example": result.get("practical_example", 0),
                "conciseness": result.get("conciseness", 0),
                "strengths": result.get("strengths", []),
                "weaknesses": result.get("weaknesses", []),
                "verdict": result.get("verdict", ""),
                "ideal_answer": result.get("ideal_answer", ""),
                "followup_question": result.get("followup_question", "")
            },
            "improvements": result.get("improvements", [])

        }

    except Exception as e:
        print("EVALUATE ANSWER ERROR")
        print(e)

        return {
            "error": str(e)
        }
    
    
@app.post("/tailor-resume")
def tailor_resume(data: ResumeTailorRequest):
    try:
        prompt = f"""
You are an expert resume writer and ATS optimization specialist.

Your task is to tailor a candidate's resume for ANY profession.

The job description may belong to:
- Software Engineering
- Data Engineering
- Data Science
- AI / ML
- HR
- Marketing
- Sales
- Finance
- Banking
- Accounting
- Healthcare
- Teaching
- Business Development
- Customer Support
- Operations
- Product Management
- Design
- Consulting
- Administration
- Government
or any other profession.

Rules:

1. Never invent experience.

2. Never invent projects.

3. Never invent certifications.

4. Never invent achievements.

5. Never add skills the candidate doesn't already have.

6. Improve wording using ATS-friendly language.

7. Rewrite bullet points using stronger action verbs.

8. Reorder skills according to the job description.

9. Improve the professional summary according to the target role.

10. Keep facts exactly the same.

11. Preserve formatting.

12. Return ONLY JSON.

Resume:

{resume_text}

Job Description:

{job_description}

Return:

{
    "target_role":"",
    "summary":{
        "original":"",
        "tailored":"",
        "reason":""
    },
    "skills":{
        "original":[],
        "tailored":[],
        "reason":""
    },
    "projects":[
        {
            "original":"",
            "tailored":"",
            "reason":""
        }
    ],
    "experience":[
        {
            "original":"",
            "tailored":"",
            "reason":""
        }
    ],
    "keywords":[],
    "missing_keywords":[],
    "suggestions":[]
}
"""

        response = None
        last_error = None
        
        for attempt in range(4):
            try:
                response = client.models.generate_content( 
                    model="gemini-3.5-flash",
                    contents=prompt,
                    config={
                        "response_mime_type": "application/json"
                    }
                )
                break
            except Exception as error:
                last_error = error
                error_text = str(error)

                if "503" in error_text or "UNAVAILABLE" in error_text:
                    if attempt < 3:
                        wait_time = (2 ** attempt) + random.uniform(0, 1)
                        time.sleep(wait_time)
                        continue
                
                raise error
        if response is None:
            raise Exception(
                "Server is temporarily busy. Please try again after a minute."
            ) from last_error


        result = json.loads(response.text)

        summary = result.get("summary", {})
        skills = result.get("skills", {})

        return {
            "target_role": result.get("target_role", ""),

            "summary": {
                "original": summary.get("original", ""),
                "tailored": summary.get("tailored", ""),
                "reason": summary.get("reason", "")
            },

            "skills": {
                "original": skills.get("original", []),
                "tailored": skills.get("tailored", []),
                "reason": skills.get("reason", "")
            },

            "projects": result.get("projects", []),
            "experience": result.get("experience", []),
            "keywords": result.get("keywords", []),
            "missing_keywords": result.get("missing_keywords", []),
            "suggestions": result.get("suggestions", [])
        }

    except json.JSONDecodeError:
        return {
            "error": "The AI response was not valid JSON. Please try again."
        }

    except Exception as error:
        return {
            "error": str(error)
        }
    
@app.post("/extract-docx")
async def extract_docx(file: UploadFile = File(...)):
    filename = file.filename or ""

    if not filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a DOCX file."
        )

    try:
        file_content = await file.read()

        if not file_content:
            raise HTTPException(
                status_code=400,
                detail="The uploaded file is empty."
            )

        document = Document(BytesIO(file_content))

        paragraphs = []

        for index, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()

            if text:
                paragraphs.append({
                    "index": index,
                    "text": text,
                    "style": paragraph.style.name
                    if paragraph.style
                    else "Normal",
                    "run_count": len(paragraph.runs),
                })

        tables = []

        for table_index, table in enumerate(document.tables):
            rows = []

            for row_index, row in enumerate(table.rows):
                cells = []

                for cell_index, cell in enumerate(row.cells):
                    cells.append({
                        "cell_index": cell_index,
                        "text": cell.text.strip(),
                    })

                rows.append({
                    "row_index": row_index,
                    "cells": cells,
                })

            tables.append({
                "table_index": table_index,
                "rows": rows,
            })

        full_text = []

        for item in paragraphs:
            full_text.append(item["text"])

        for table in tables:
            for row in table["rows"]:
                for cell in row["cells"]:
                    if cell["text"]:
                        full_text.append(cell["text"])

        return {
            "filename": filename,
            "paragraph_count": len(document.paragraphs),
            "non_empty_paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
            "paragraphs": paragraphs,
            "tables": tables,
            "full_text": "\n".join(full_text),
        }

    except HTTPException:
        raise

    except Exception as error:
        print("DOCX extraction error:", error)

        raise HTTPException(
            status_code=500,
            detail="The DOCX file could not be read."
        )
    
@app.post("/resume-health")
async def resume_health(data: ResumeHealthRequest):

    prompt = f"""
You are an expert ATS Resume Reviewer.

Evaluate this resume for ANY profession.

The resume may belong to Software Engineering, HR, Marketing,
Finance, Teaching, Healthcare, Design, Sales, Government,
Business or any other profession.

Resume:

{data.resume_text}

ATS Score:
{data.ats_score}

Matched Skills:
{", ".join(data.matched_skills)}

Missing Skills:
{", ".join(data.missing_skills)}

Evaluate objectively.

Do NOT assume the candidate is from Software Engineering.

Judge the resume according to the standards of that profession.

For example:

Software → projects, coding, technologies

Teacher → teaching experience, classroom management

HR → recruitment, onboarding, policies

Sales → targets, revenue, negotiation

Marketing → campaigns, analytics

Finance → accounting, auditing

Healthcare → patient care, certifications

Business → operations, management

etc.

1. ATS Compatibility
2. Professional Summary
3. Experience Quality
4. Skills Relevance
5. Grammar
6. Action Verbs
7. Achievements
8. Overall Resume Quality

Rules:

• Do not assume projects exist.
• Do not assume technical roles.
• Score each category from 0-100.
• Give practical feedback.
• Never invent information.
• If a section does not exist (for example projects or experience), score based only on the information actually present.
• Do not penalize candidates simply because they belong to another profession.

Return ONLY JSON.

{
  "overall":0,
  "ats":0,
  "summary":0,
  "experience":0,
  "skills":0,
  "grammar":0,
  "action_verbs":0,
  "achievements":0,
  "feedback":[]
}
"""

    try:

        response = client.models.generate_content(
            model="gemini-3.5-flash",
            contents=prompt,
            config={
                "response_mime_type": "application/json"
            }
        )

        result = json.loads(response.text)

        return result

    except Exception as e:

        print("Resume Health Error:", e)

        overall = max(60, data.ats_score)

        return {
            "overall": overall,
            "ats": data.ats_score,
            "summary": 80,
            "experience": 80,
            "skills": 80,
            "grammar": 90,
            "action_verbs": 80,
            "achievements": 75,
            "feedback": [
                "AI evaluation is temporarily unavailable.",
                "Your ATS score has still been calculated successfully."
            ]
        }
    
    
@app.post("/preserve-docx-test")
async def preserve_docx_test(file: UploadFile = File(...)):
    filename = file.filename or ""

    if not filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a DOCX file.",
        )

    content = await file.read()

    if not content:
        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty.",
        )

    try:
        document = Document(BytesIO(content))

        output = BytesIO()
        document.save(output)
        output.seek(0)

        output_name = f"preserved_{Path(filename).name}"

        return StreamingResponse(
            output,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": (
                    f'attachment; filename="{output_name}"'
                )
            },
        )

    except Exception as error:
        print("DOCX preservation error:", error)

        raise HTTPException(
            status_code=500,
            detail="The DOCX file could not be processed.",
        ) 
@app.post("/replace-docx-text-test")
async def replace_docx_text_test(
    file: Annotated[UploadFile, File()],
    location_id: Annotated[str, Form()],
    replacement_text: Annotated[str, Form()],
):
    filename = file.filename or ""

    if not filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a DOCX file.",
        )

    file_content = await file.read()

    if not file_content:
        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty.",
        )

    try:
        document = Document(BytesIO(file_content))
        locations = collect_docx_text_locations(document)

        selected_location = next(
            (
                location
                for location in locations
                if location["location_id"] == location_id
            ),
            None,
        )

        if not selected_location:
            raise HTTPException(
                status_code=404,
                detail="The selected text location was not found.",
            )

        paragraph = get_docx_paragraph_by_location(
            document,
            selected_location,
        )

        replace_paragraph_text_preserving_style(
            paragraph,
            replacement_text,
        )

        output = BytesIO()
        document.save(output)
        output.seek(0)

        original_name = Path(filename).stem
        output_name = f"{original_name}_replacement_test.docx"

        return StreamingResponse(
            output,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": (
                    f'attachment; filename="{output_name}"'
                )
            },
        )

    except HTTPException:
        raise

    except Exception as error:
        print("DOCX replacement test error:", error)

        raise HTTPException(
            status_code=500,
            detail="The DOCX text could not be replaced.",
        )
@app.post("/optimize-docx")
async def optimize_docx(
    file: Annotated[UploadFile, File()],
    job_description: Annotated[str, Form()],
):
    filename = file.filename or ""

    if not filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a DOCX file.",
        )

    if not job_description.strip():
        raise HTTPException(
            status_code=400,
            detail="Job description is required.",
        )

    file_content = await file.read()

    if not file_content:
        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty.",
        )

    try:
        document = Document(BytesIO(file_content))
        locations = collect_docx_text_locations(document)

        editable_locations = []

        current_section = ""

        for location in locations:
            text = location["text"].strip()

            if not text:
                continue
            upper = text.upper()

            if upper in {
                "SUMMARY",
                "EDUCATION",
                "PROJECTS",
                "SKILLS",
                "CERTIFICATIONS",
                "CODING PROFILES",
            }:
                current_section = upper
                continue
            if current_section not in {
                "SUMMARY",
                "PROJECTS",
                "SKILLS",
            }:
                continue

            if len(text) < 20:
                continue

            editable_locations.append({
                "location_id": location["location_id"],
                "section": current_section,
                "text": text,
                "max_characters": int(len(text) * 1.25),
            })

        prompt = f"""
You are an expert ATS resume editor.

Improve the resume text for the supplied job description.

STRICT RULES:
- Preserve every existing fact.
- Do not invent skills, achievements, metrics, projects, experience, or certifications.
- Keep the same meaning.
- Improve clarity, impact, action verbs, and job relevance.
- Keep each improved text close to the original length.
- Do not change headings.
- Do not add new sections.
- If a sentence is already strong, keep it nearly unchanged.
- Return only valid JSON.

JOB DESCRIPTION:
{job_description[:6000]}

TEXT LOCATIONS:
{json.dumps(editable_locations, ensure_ascii=False)}

Return exactly:

{{
  "changes": [
    {{
      "location_id": "paragraph:5",
      "original": "original text",
      "improved": "improved text"
    }}
  ]
}}
"""
        response = None
        last_error = None
        model_candidates = [
            "gemini-3.5-flash",
            "gemini-2.5-flash",
        ]

        for model_name in model_candidates:
            for attempt in range(4):
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt,
                        config={
                            "response_mime_type": "application/json"
                        }
                    )
                    break
                except Exception as error:
                    last_error = error
                    error_text = str(error)
                is_temporary = (
                    "503" in error_text or "UNAVAILABLE" in error_text or "RESOURCE_EXHAUSTED" in error_text
                )    
                if is_temporary and attempt < 3:
                    wait_seconds = (2 ** attempt) + random.uniform(0, 1)
                    time.sleep(wait_seconds)
                    continue
                break
                
            if response is None:
               break
        if response is None:

           raise HTTPException(
                status_code=503,
                detail=(
                   "Gemini is temporarily unavailable after trying multiple "
                   "models. Please try again in a few minutes." 
                )
            ) from last_error
                
        
        result = json.loads(response.text)
        changes = result.get("changes", [])

        location_map = {
            location["location_id"]: location
            for location in locations
        }

        applied_changes = []

        for change in changes:
            location_id = change.get("location_id", "")
            improved_text = change.get("improved", "").strip()

            if not location_id or not improved_text:
                continue

            location = location_map.get(location_id)

            if not location:
                continue

            original_text = location["text"].strip()

            if original_text != change.get("original", "").strip():
                continue

            paragraph = get_docx_paragraph_by_location(
                document,
                location,
            )

            replace_paragraph_text_preserving_style(
                paragraph,
                improved_text,
            )

            applied_changes.append({
                "location_id": location_id,
                "original": original_text,
                "improved": improved_text,
            })

        output = BytesIO()
        document.save(output)
        output.seek(0)

        original_name = Path(filename).stem
        output_name = f"{original_name}_optimized.docx"

        return StreamingResponse(
            output,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": (
                    f'attachment; filename="{output_name}"'
                ),
                "X-Applied-Changes": str(len(applied_changes)),
            },
        )

    except HTTPException:
        raise

    except json.JSONDecodeError:
        raise HTTPException(
            status_code=502,
            detail="The AI response was not valid JSON.",
        )

    except Exception as error:
        print("DOCX optimization error:", error)

        raise HTTPException(
            status_code=500,
            detail=str(error),
        )
@app.post("/extract-pdf-layout")
async def extract_pdf_layout(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF file."
        )

    pdf_bytes = await file.read()

    try:
        document = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages = []

        for page_number, page in enumerate(document):
            page_dict = page.get_text("dict")
            page_blocks = []

            for block_index, block in enumerate(page_dict["blocks"]):
                if "lines" not in block:
                    continue

                block_text = []
                first_span = None

                for line in block["lines"]:
                    for span in line["spans"]:
                        if first_span is None:
                            first_span = span

                        block_text.append(span["text"])

                cleaned_text = " ".join(block_text).strip()

                if not cleaned_text or first_span is None:
                    continue

                x0, y0, x1, y1 = block["bbox"]

                page_blocks.append({
                    "block_id": f"page:{page_number}:block:{block_index}",
                    "text": cleaned_text,
                    "x0": x0,
                    "y0": y0,
                    "x1": x1,
                    "y1": y1,
                    "width": x1 - x0,
                    "height": y1 - y0,
                    "font": first_span["font"],
                    "font_size": first_span["size"],
                    "color": first_span["color"],
                })

            # Keep this outside the block loop
            pages.append({
                "page_number": page_number + 1,
                "width": page.rect.width,
                "height": page.rect.height,
                "blocks": page_blocks,
            })

        total_pages = len(pages)
        document.close()

        return {
            "filename": file.filename,
            "total_pages": total_pages,
            "pages": pages,
        }

    except Exception as error:
        raise HTTPException(
            status_code=500,
            detail=f"Could not extract PDF layout: {str(error)}"
        )
@app.post("/replace-pdf-text-test")
async def replace_pdf_text_test(
    file: UploadFile = File(...)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF file."
        )

    pdf_bytes = await file.read()

    try:
        document = fitz.open(stream=pdf_bytes, filetype="pdf")
        page = document[0]

        search_text = "AI & ML Engineer"
        replacement_text = "Software Engineer"

        locations = page.search_for(search_text)

        if not locations:
            document.close()
            raise HTTPException(
                status_code=404,
                detail="Text was not found in the PDF."
            )

        rect = locations[0]

        # Actually remove the original text
        page.add_redact_annot(
            rect,
            fill=(1, 1, 1)
        )

        page.apply_redactions()

        # Add the replacement text in the same position

        replacement_rect = fitz.Rect(
           rect.x0,
           rect.y0 - 3,
           rect.x1 + 100,
           rect.y1 + 5
        )

        result = page.insert_textbox(
            replacement_rect,
            replacement_text,
            fontsize=9,
            fontname="helv",
            color=(0, 0, 0),
            align=fitz.TEXT_ALIGN_LEFT
        )

        print("Textbox result:", result)


        output = BytesIO()
        document.save(output)
        output.seek(0)
        document.close()

        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={
                "Content-Disposition":
                "attachment; filename=updated_resume.pdf"
            }
        )

    except HTTPException:
        raise

    except Exception as error:
        raise HTTPException(
            status_code=500,
            detail=f"Could not replace PDF text: {str(error)}"
        )
@app.post("/optimize-pdf")
async def optimize_pdf(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400,
            detail="Please upload a PDF file."
        )

    if not job_description.strip():
        raise HTTPException(
            status_code=400,
            detail="Please provide a job description."
        )

    pdf_bytes = await file.read()

    if not pdf_bytes:
        raise HTTPException(
            status_code=400,
            detail="The uploaded PDF is empty."
        )

    document = None

    try:
        document = fitz.open(
            stream=pdf_bytes,
            filetype="pdf"
        )

        editable_blocks = []
        detected_headings = []
        current_section = ""

        heading_aliases = {
            "SUMMARY": "SUMMARY",
            "PROFILE": "SUMMARY",
            "PROFESSIONALSUMMARY": "SUMMARY",
            "CAREERSUMMARY": "SUMMARY",
            "OBJECTIVE": "SUMMARY",
            "CAREEROBJECTIVE": "SUMMARY",
            "ABOUTME": "SUMMARY",

            "PROJECTS": "PROJECTS",
            "PROJECT": "PROJECTS",
            "PERSONALPROJECTS": "PROJECTS",
            "ACADEMICPROJECTS": "PROJECTS",
            "KEYPROJECTS": "PROJECTS",
            "PROJECTEXPERIENCE": "PROJECTS",

            "SKILLS": "SKILLS",
            "TECHNICALSKILLS": "SKILLS",
            "CORESKILLS": "SKILLS",
            "KEYSKILLS": "SKILLS",
            "SKILLSET": "SKILLS",
            "TECHNOLOGIES": "SKILLS",
            "TECHNOLOGYSTACK": "SKILLS",

            "EDUCATION": "EDUCATION",
            "ACADEMICBACKGROUND": "EDUCATION",
            "ACADEMICQUALIFICATIONS": "EDUCATION",

            "EXPERIENCE": "EXPERIENCE",
            "WORKEXPERIENCE": "EXPERIENCE",
            "PROFESSIONALEXPERIENCE": "EXPERIENCE",
            "INTERNSHIP": "EXPERIENCE",
            "INTERNSHIPS": "EXPERIENCE",

            "CERTIFICATIONS": "CERTIFICATIONS",
            "CERTIFICATES": "CERTIFICATIONS",
            "COURSES": "CERTIFICATIONS",

            "CODINGPROFILES": "CODING PROFILES",
            "CODINGPROFILE": "CODING PROFILES",
            "PROFILES": "CODING PROFILES",

            "ACHIEVEMENTS": "ACHIEVEMENTS",
            "AWARDS": "ACHIEVEMENTS",

            "POSITIONSOFRESPONSIBILITY":
                "POSITIONS OF RESPONSIBILITY",

            "CONTACT": "CONTACT",
            "CONTACTDETAILS": "CONTACT",
        }

        allowed_sections = {
            "SUMMARY",
            "PROJECTS",
            "SKILLS",
        }

        skill_labels = [
            "Languages:",
            "Web Technologies:",
            "Tools & Platforms:",
            "Database:",
            "AI/ML:",
            "Cloud & APIs:",
        ]

        def clean_text(value: str) -> str:
            return " ".join(str(value).split()).strip()

        def normalize_heading(value: str) -> str:
            normalized = clean_text(value).upper()

            normalized = normalized.rstrip(
                ":|-–—"
            ).strip()

            # Handles headings extracted like:
            # S UMMARY, P ROJECTS, S KILLS
            normalized = normalized.replace(" ", "")

            return normalized

        def detect_heading(value: str):
            normalized = normalize_heading(value)

            return heading_aliases.get(normalized)

        def replace_unsupported_characters(
            value: str
        ) -> str:
            replacements = {
                "•": "-",
                "●": "-",
                "▪": "-",
                "◦": "-",
                "–": "-",
                "—": "-",
                "−": "-",
                "’": "'",
                "‘": "'",
                "“": '"',
                "”": '"',
                "\u00a0": " ",
            }

            for old_character, new_character in replacements.items():
                value = value.replace(
                    old_character,
                    new_character
                )

            return value

        def safe_pdf_text(value: str) -> str:
            value = replace_unsupported_characters(
                str(value)
            )

            return clean_text(value)

        def format_skills_text(value: str) -> str:
            value = replace_unsupported_characters(
                str(value)
            )

            # Remove existing line breaks first.
            value = clean_text(value)

            # Correct the known typo from the original resume.
            value = value.replace(
                "EST APIs",
                "REST APIs"
            )

            for label in skill_labels:
                value = value.replace(
                    label,
                    "\n" + label
                )

            lines = []

            for line in value.splitlines():
                cleaned_line = clean_text(line)

                if cleaned_line:
                    lines.append(cleaned_line)

            return "\n".join(lines).strip()

        def shorten_at_word(
            value: str,
            maximum_length: int
        ) -> str:
            if len(value) <= maximum_length:
                return value

            shortened = value[:maximum_length]

            if " " in shortened:
                shortened = shortened.rsplit(
                    " ",
                    1
                )[0]

            return shortened.rstrip(
                " ,;:-"
            )

        print("\n" + "=" * 70)
        print("PDF OPTIMIZATION")
        print("=" * 70)

        for page_number, page in enumerate(document):
            page_dict = page.get_text("dict")

            for block_index, block in enumerate(
                page_dict.get("blocks", [])
            ):
                lines = block.get("lines", [])

                if not lines:
                    continue

                block_spans = []
                block_lines = []

                for line in lines:
                    spans = line.get("spans", [])

                    if not spans:
                        continue

                    block_spans.extend(spans)

                    line_text = clean_text(
                        " ".join(
                            span.get("text", "")
                            for span in spans
                            if span.get("text", "").strip()
                        )
                    )

                    if line_text:
                        block_lines.append(line_text)

                if not block_lines or not block_spans:
                    continue

                block_text = clean_text(
                    " ".join(block_lines)
                )

                if not block_text:
                    continue

                matched_section = detect_heading(
                    block_text
                )

                if matched_section:
                    current_section = matched_section

                    if matched_section not in detected_headings:
                        detected_headings.append(
                            matched_section
                        )

                    print(
                        f"Detected section: {matched_section}"
                    )

                    # Never edit section headings.
                    continue

                if current_section not in allowed_sections:
                    continue

                should_edit = False

                if current_section == "SUMMARY":
                    should_edit = True

                elif current_section == "SKILLS":
                    should_edit = True

                elif current_section == "PROJECTS":
                    # Only edit project bullet descriptions.
                    # Keep project names, roles, dates and
                    # technology lines unchanged.
                    first_character = block_text.lstrip()[:1]

                    bullet_characters = {
                        "•",
                        "●",
                        "▪",
                        "◦",
                        "-",
                        "–",
                        "—",
                    }

                    if first_character in bullet_characters:
                        should_edit = True

                if not should_edit:
                    continue

                x0, y0, x1, y1 = block["bbox"]
                first_span = block_spans[0]

                editable_blocks.append({
                    "block_id": (
                        f"page:{page_number}:"
                        f"block:{block_index}"
                    ),
                    "page_number": page_number,
                    "section": current_section,
                    "text": block_text,
                    "x0": float(x0),
                    "y0": float(y0),
                    "x1": float(x1),
                    "y1": float(y1),
                    "font_size": float(
                        first_span.get("size", 10)
                    ),
                    "max_characters": max(
                        len(block_text),
                        20
                    ),
                })

        print(
            "Detected headings:",
            detected_headings
        )

        print(
            "Editable blocks found:",
            len(editable_blocks)
        )

        print("=" * 70 + "\n")

        if not editable_blocks:
            detected_text = (
                ", ".join(detected_headings)
                if detected_headings
                else "None"
            )

            raise HTTPException(
                status_code=400,
                detail=(
                    "No editable resume sections were found. "
                    f"Detected headings: {detected_text}."
                )
            )

        prompt = f"""
You are optimizing selected resume content for the supplied job description.

JOB DESCRIPTION:

{job_description}

RESUME BLOCKS:

{json.dumps(editable_blocks, ensure_ascii=False)}

Rules:

1. Preserve every candidate fact.
2. Do not invent experience, skills, tools, technologies,
   responsibilities, metrics, qualifications or achievements.
3. Do not change project names.
4. Do not change company names.
5. Do not change job titles or project roles.
6. Do not change dates, numbers, URLs or certifications.
7. Do not add any skill absent from the original resume block.
8. Improve SUMMARY for alignment with the job description.
9. Improve PROJECT descriptions using concise action-oriented language.
10. Reorder SKILLS according to relevance using only existing skills.
11. Preserve all skill category labels exactly.
12. Keep each skill category separate.
13. Preserve the meaning of every original sentence.
14. Do not exceed max_characters for SUMMARY or PROJECT blocks.
15. Keep every block_id exactly unchanged.
16. Return one result for every supplied block.
17. Return valid JSON only.
18. Do not use Markdown.
19. Do not provide explanations.
20. Use standard ASCII hyphens instead of Unicode bullet symbols.
21. Do not leave incomplete or unfinished sentences.
22. For project descriptions, keep a leading hyphen.
23. Do not rename Website to System, Platform or any other term.

Return exactly this JSON structure:

{{
  "changes": [
    {{
      "block_id": "page:0:block:4",
      "improved": "Improved text"
    }}
  ]
}}
"""

        response = None
        last_error = None

        model_candidates = [
            "gemini-3.5-flash",
            "gemini-2.5-flash",
        ]

        for model_name in model_candidates:
            for attempt in range(4):
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=prompt,
                        config={
                            "response_mime_type":
                                "application/json"
                        }
                    )

                    if response and response.text:
                        break

                except Exception as error:
                    last_error = error
                    error_text = str(error).upper()

                    temporary_error = (
                        "503" in error_text
                        or "UNAVAILABLE" in error_text
                        or "429" in error_text
                        or "RESOURCE_EXHAUSTED"
                        in error_text
                    )

                    if temporary_error and attempt < 3:
                        wait_seconds = (
                            2 ** attempt
                        ) + random.uniform(0, 1)

                        time.sleep(wait_seconds)
                        continue

                    break

            if response and response.text:
                break

        if not response or not response.text:
            raise HTTPException(
                status_code=503,
                detail=(
                    "Gemini is temporarily unavailable. "
                    "Please try again after a few minutes."
                )
            ) from last_error

        try:
            result = json.loads(response.text)

        except json.JSONDecodeError as error:
            print("Invalid Gemini response:")
            print(response.text)

            raise HTTPException(
                status_code=500,
                detail="Gemini returned invalid JSON."
            ) from error

        changes = result.get("changes", [])

        if not isinstance(changes, list):
            raise HTTPException(
                status_code=500,
                detail=(
                    "Gemini returned an invalid changes format."
                )
            )

        blocks_by_id = {
            block["block_id"]: block
            for block in editable_blocks
        }

        valid_changes = []

        for change in changes:
            if not isinstance(change, dict):
                continue

            block_id = change.get("block_id")

            original_block = blocks_by_id.get(
                block_id
            )

            if not original_block:
                continue

            raw_improved_text = change.get(
                "improved",
                ""
            )

            if original_block["section"] == "SKILLS":
                improved_text = format_skills_text(
                    raw_improved_text
                )

            else:
                improved_text = safe_pdf_text(
                    raw_improved_text
                )

            if not improved_text:
                continue

            max_characters = original_block[
                "max_characters"
            ]

            # Do not cut Skills because its line breaks
            # require a slightly different text length.
            if original_block["section"] != "SKILLS":
                improved_text = shorten_at_word(
                    improved_text,
                    max_characters
                )

            if original_block["section"] == "PROJECTS":
                improved_text = improved_text.lstrip(
                    "- "
                )

                improved_text = "- " + improved_text

                improved_text = shorten_at_word(
                    improved_text,
                    max_characters
                )

            valid_changes.append({
                "block_id": block_id,
                "improved": improved_text,
            })

        if not valid_changes:
            raise HTTPException(
                status_code=500,
                detail=(
                    "Gemini did not return any valid "
                    "resume changes."
                )
            )

        # First pass: remove original editable text.
        for change in valid_changes:
            original_block = blocks_by_id[
                change["block_id"]
            ]

            page = document[
                original_block["page_number"]
            ]

            original_rect = fitz.Rect(
                original_block["x0"] - 1,
                original_block["y0"] - 1,
                original_block["x1"] + 2,
                original_block["y1"] + 2,
            )

            page.add_redact_annot(
                original_rect,
                fill=(1, 1, 1)
            )

        for page in document:
            page.apply_redactions()

        # Second pass: insert optimized text.
        for change in valid_changes:
            original_block = blocks_by_id[
                change["block_id"]
            ]

            improved_text = change["improved"]

            page = document[
                original_block["page_number"]
            ]

            original_font_size = original_block[
                "font_size"
            ]

            page_width = page.rect.width

            if original_block["section"] == "SUMMARY":
                extra_height = 14
                extra_width = 80

            elif original_block["section"] == "SKILLS":
                extra_height = 8
                extra_width = 40

            else:
                extra_height = 8
                extra_width = 40

            replacement_rect = fitz.Rect(
                original_block["x0"],
                original_block["y0"] - 1,
                min(
                    original_block["x1"] + extra_width,
                    page_width - 20
                ),
                original_block["y1"] + extra_height,
            )

            font_sizes = [
                original_font_size,
                original_font_size - 0.5,
                original_font_size - 1,
                original_font_size - 1.5,
                original_font_size - 2,
            ]

            inserted = False

            for font_size in font_sizes:
                safe_font_size = max(
                    float(font_size),
                    7.0
                )

                insert_result = page.insert_textbox(
                    replacement_rect,
                    improved_text,
                    fontsize=safe_font_size,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

                if insert_result >= 0:
                    inserted = True
                    break

            if not inserted:
                larger_rect = fitz.Rect(
                    replacement_rect.x0,
                    replacement_rect.y0,
                    replacement_rect.x1,
                    replacement_rect.y1 + 15,
                )

                final_result = page.insert_textbox(
                    larger_rect,
                    improved_text,
                    fontsize=7.0,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

                if final_result < 0:
                    print(
                        "Warning: text could not fit for",
                        change["block_id"]
                    )

        output = BytesIO()

        document.save(
            output,
            garbage=4,
            deflate=True
        )

        output.seek(0)

        document.close()
        document = None

        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={
                "Content-Disposition": (
                    'attachment; '
                    'filename="optimized_resume.pdf"'
                )
            }
        )

    except HTTPException:
        if document is not None:
            document.close()

        raise

    except Exception as error:
        if document is not None:
            document.close()

        print(
            "PDF optimization error:",
            error
        )

        raise HTTPException(
            status_code=500,
            detail=(
                f"Could not optimize PDF: {str(error)}"
            )
        )
    